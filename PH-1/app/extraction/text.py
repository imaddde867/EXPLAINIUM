import pdfplumber
import docx
from fastapi import UploadFile, HTTPException
from typing import Optional, List, Dict
import pytesseract
from PIL import Image
import io
import cv2
import numpy as np
import tempfile
import os
import logging

logger = logging.getLogger(__name__)

def extract_text_pdf(file: UploadFile) -> Optional[str]:
    """Extract text from PDF files"""
    try:
        file.file.seek(0)
        with pdfplumber.open(file.file) as pdf:
            text_content = []
            for page_num, page in enumerate(pdf.pages):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text_content.append(f"--- Page {page_num + 1} ---\n{page_text}")
                except Exception as e:
                    logger.warning(f"Failed to extract text from page {page_num + 1}: {e}")
                    continue
            return "\n\n".join(text_content) if text_content else None
    except Exception as e:
        logger.error(f"PDF extraction failed: {e}")
        raise HTTPException(status_code=500, detail=f"PDF processing failed: {str(e)}")

def extract_text_docx(file: UploadFile) -> Optional[str]:
    """Extract text from DOCX files"""
    try:
        file.file.seek(0)
        doc = docx.Document(file.file)
        text_content = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_content.append(para.text)
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_content.append(" | ".join(row_text))
        return "\n".join(text_content) if text_content else None
    except Exception as e:
        logger.error(f"DOCX extraction failed: {e}")
        raise HTTPException(status_code=500, detail=f"DOCX processing failed: {str(e)}")

def extract_text_txt(file: UploadFile) -> Optional[str]:
    """Extract text from plain text files"""
    try:
        file.file.seek(0)
        content = file.file.read()
        # Try different encodings
        encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
        for encoding in encodings:
            try:
                return content.decode(encoding)
            except UnicodeDecodeError:
                continue
        # If all encodings fail, use utf-8 with error handling
        return content.decode('utf-8', errors='replace')
    except Exception as e:
        logger.error(f"Text extraction failed: {e}")
        raise HTTPException(status_code=500, detail=f"Text processing failed: {str(e)}")

def extract_text_image(file: UploadFile) -> Optional[str]:
    """Extract text from images using OCR with preprocessing"""
    try:
        # Import here to avoid circular imports
        from .image import extract_text_from_image
        return extract_text_from_image(file, preprocess=True)
    except Exception as e:
        logger.error(f"OCR extraction failed: {e}")
        raise HTTPException(status_code=500, detail=f"Image OCR processing failed: {str(e)}")

def extract_video_keyframes(file: UploadFile, frame_interval: int = 30, max_frames: int = 50) -> List[Dict]:
    """Extract key frames from video files"""
    try:
        # Import here to avoid circular imports
        from .video import extract_video_keyframes as extract_frames
        return extract_frames(file, frame_interval, max_frames)
    except Exception as e:
        logger.error(f"Video frame extraction failed: {e}")
        raise HTTPException(status_code=500, detail=f"Video processing failed: {str(e)}")