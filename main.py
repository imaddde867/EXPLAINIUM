#!/usr/bin/env python3
"""
Universal Document Preprocessor for Factory/Technical Documentation
Handles diverse file formats and extracts structured data for knowledge base ingestion.
"""

import os
import logging
import mimetypes
import hashlib
import json
from pathlib import Path
from typing import Dict, List, Any, Optional, Tuple
from dataclasses import dataclass, asdict
from datetime import datetime
import asyncio
import aiofiles
from concurrent.futures import ThreadPoolExecutor, ProcessPoolExecutor

# Document Processing Libraries
import PyPDF2
import pdfplumber
from docx import Document
import pandas as pd
import openpyxl
from PIL import Image
import pytesseract
import cv2
import numpy as np

# Media Processing
import moviepy.editor as mp
import librosa
import speech_recognition as sr

# NLP Processing
import spacy
from transformers import pipeline, AutoTokenizer, AutoModel
import torch
from sentence_transformers import SentenceTransformer

# Storage and Database
import sqlite3
import json
from sqlalchemy import create_engine, Column, String, Text, DateTime, Integer, LargeBinary
from sqlalchemy.ext.declarative import declarative_base
from sqlalchemy.orm import sessionmaker

# Configuration
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

@dataclass
class ProcessedDocument:
    """Container for processed document data"""
    file_path: str
    file_type: str
    file_hash: str
    extracted_text: str
    metadata: Dict[str, Any]
    entities: List[Dict[str, Any]]
    embeddings: Optional[List[float]]
    processing_timestamp: datetime
    confidence_score: float
    extracted_media: List[Dict[str, Any]]

class FormatDetector:
    """Intelligent file format detection and classification"""
    
    SUPPORTED_FORMATS = {
        'documents': ['.pdf', '.docx', '.doc', '.txt', '.rtf', '.odt'],
        'spreadsheets': ['.xlsx', '.xls', '.csv', '.ods'],
        'images': ['.jpg', '.jpeg', '.png', '.bmp', '.tiff', '.gif'],
        'videos': ['.mp4', '.avi', '.mov', '.wmv', '.flv', '.mkv'],
        'audio': ['.mp3', '.wav', '.m4a', '.flac', '.aac'],
        'archives': ['.zip', '.rar', '.7z', '.tar', '.gz'],
        'cad': ['.dwg', '.dxf', '.step', '.stp', '.iges'],
        'presentations': ['.pptx', '.ppt', '.odp']
    }
    
    @classmethod
    def detect_format(cls, file_path: str) -> Tuple[str, str]:
        """Detect file format and category"""
        file_ext = Path(file_path).suffix.lower()
        mime_type, _ = mimetypes.guess_type(file_path)
        
        for category, extensions in cls.SUPPORTED_FORMATS.items():
            if file_ext in extensions:
                return category, file_ext
        
        return 'unknown', file_ext
    
    @classmethod
    def is_supported(cls, file_path: str) -> bool:
        """Check if file format is supported"""
        category, _ = cls.detect_format(file_path)
        return category != 'unknown'

class TextExtractor:
    """Handles text extraction from various document formats"""
    
    def __init__(self):
        self.ocr_config = '--oem 3 --psm 6'
        
    def extract_from_pdf(self, file_path: str) -> Dict[str, Any]:
        """Extract text and metadata from PDF files"""
        try:
            # Try pdfplumber first for better table extraction
            with pdfplumber.open(file_path) as pdf:
                text_content = []
                tables = []
                images = []
                
                for page_num, page in enumerate(pdf.pages):
                    # Extract text
                    page_text = page.extract_text()
                    if page_text:
                        text_content.append(f"--- Page {page_num + 1} ---\n{page_text}")
                    
                    # Extract tables
                    page_tables = page.extract_tables()
                    if page_tables:
                        for table_idx, table in enumerate(page_tables):
                            tables.append({
                                'page': page_num + 1,
                                'table_id': table_idx,
                                'data': table
                            })
                
                return {
                    'text': '\n\n'.join(text_content),
                    'tables': tables,
                    'images': images,
                    'page_count': len(pdf.pages),
                    'metadata': pdf.metadata or {}
                }
                
        except Exception as e:
            logger.warning(f"pdfplumber failed for {file_path}: {e}")
            # Fallback to PyPDF2
            return self._extract_pdf_fallback(file_path)
    
    def _extract_pdf_fallback(self, file_path: str) -> Dict[str, Any]:
        """Fallback PDF extraction using PyPDF2"""
        try:
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                text_content = []
                
                for page_num, page in enumerate(reader.pages):
                    text_content.append(f"--- Page {page_num + 1} ---\n{page.extract_text()}")
                
                return {
                    'text': '\n\n'.join(text_content),
                    'tables': [],
                    'images': [],
                    'page_count': len(reader.pages),
                    'metadata': reader.metadata or {}
                }
        except Exception as e:
            logger.error(f"PDF extraction failed: {e}")
            return {'text': '', 'tables': [], 'images': [], 'metadata': {}}
    
    def extract_from_docx(self, file_path: str) -> Dict[str, Any]:
        """Extract text and metadata from Word documents"""
        try:
            doc = Document(file_path)
            
            # Extract paragraphs
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            
            # Extract tables
            tables = []
            for table_idx, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                tables.append({'table_id': table_idx, 'data': table_data})
            
            return {
                'text': '\n\n'.join(paragraphs),
                'tables': tables,
                'images': [],
                'metadata': dict(doc.core_properties.__dict__) if hasattr(doc, 'core_properties') else {}
            }
            
        except Exception as e:
            logger.error(f"DOCX extraction failed: {e}")
            return {'text': '', 'tables': [], 'images': [], 'metadata': {}}
    
    def extract_from_image(self, file_path: str) -> Dict[str, Any]:
        """Extract text from images using OCR"""
        try:
            # Load and preprocess image
            image = cv2.imread(file_path)
            gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
            
            # Apply image preprocessing for better OCR
            denoised = cv2.fastNlMeansDenoising(gray)
            thresh = cv2.threshold(denoised, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)[1]
            
            # Extract text using Tesseract
            text = pytesseract.image_to_string(thresh, config=self.ocr_config)
            
            # Get detailed OCR data
            ocr_data = pytesseract.image_to_data(thresh, output_type=pytesseract.Output.DICT)
            
            return {
                'text': text,
                'ocr_confidence': np.mean([int(conf) for conf in ocr_data['conf'] if int(conf) > 0]),
                'word_count': len([w for w in ocr_data['text'] if w.strip()]),
                'metadata': {
                    'image_size': image.shape,
                    'processing_method': 'tesseract_ocr'
                }
            }
            
        except Exception as e:
            logger.error(f"Image OCR failed: {e}")
            return {'text': '', 'ocr_confidence': 0, 'metadata': {}}

class MediaProcessor:
    """Handles audio/video processing and transcription"""
    
    def __init__(self):
        self.recognizer = sr.Recognizer()
        
    def extract_from_video(self, file_path: str) -> Dict[str, Any]:
        """Extract audio and transcribe from video files"""
        try:
            # Load video
            video = mp.VideoFileClip(file_path)
            
            # Extract audio
            audio_path = f"/tmp/{Path(file_path).stem}_audio.wav"
            video.audio.write_audiofile(audio_path, verbose=False, logger=None)
            
            # Transcribe audio
            transcription = self._transcribe_audio(audio_path)
            
            # Clean up temporary file
            os.remove(audio_path)
            
            return {
                'text': transcription['text'],
                'duration': video.duration,
                'fps': video.fps,
                'resolution': video.size,
                'transcription_confidence': transcription['confidence'],
                'metadata': {
                    'format': 'video',
                    'processing_method': 'moviepy_speech_recognition'
                }
            }
            
        except Exception as e:
            logger.error(f"Video processing failed: {e}")
            return {'text': '', 'metadata': {}}
    
    def extract_from_audio(self, file_path: str) -> Dict[str, Any]:
        """Extract and transcribe audio files"""
        try:
            transcription = self._transcribe_audio(file_path)
            
            # Get audio features using librosa
            y, sr = librosa.load(file_path)
            duration = librosa.get_duration(y=y, sr=sr)
            
            return {
                'text': transcription['text'],
                'duration': duration,
                'sample_rate': sr,
                'transcription_confidence': transcription['confidence'],
                'metadata': {
                    'format': 'audio',
                    'processing_method': 'speech_recognition'
                }
            }
            
        except Exception as e:
            logger.error(f"Audio processing failed: {e}")
            return {'text': '', 'metadata': {}}
    
    def _transcribe_audio(self, audio_path: str) -> Dict[str, Any]:
        """Transcribe audio using speech recognition"""
        try:
            with sr.AudioFile(audio_path) as source:
                audio = self.recognizer.record(source)
                text = self.recognizer.recognize_google(audio)
                return {'text': text, 'confidence': 0.8}  # Google API doesn't return confidence
                
        except sr.UnknownValueError:
            return {'text': '', 'confidence': 0.0}
        except Exception as e:
            logger.error(f"Audio transcription failed: {e}")
            return {'text': '', 'confidence': 0.0}

class NLPProcessor:
    """Advanced NLP processing for extracted text"""
    
    def __init__(self):
        try:
            self.nlp = spacy.load("en_core_web_sm")
        except OSError:
            logger.warning("spaCy model not found. Install with: python -m spacy download en_core_web_sm")
            self.nlp = None
        
        # Initialize sentence transformer for embeddings
        try:
            self.sentence_model = SentenceTransformer('all-MiniLM-L6-v2')
        except Exception as e:
            logger.warning(f"Sentence transformer initialization failed: {e}")
            self.sentence_model = None
    
    def process_text(self, text: str, doc_type: str = 'technical') -> Dict[str, Any]:
        """Process text and extract structured information"""
        if not text.strip():
            return {'entities': [], 'embeddings': [], 'summary': '', 'confidence': 0.0}
        
        try:
            # Named Entity Recognition
            entities = self._extract_entities(text)
            
            # Generate embeddings
            embeddings = self._generate_embeddings(text)
            
            # Extract technical patterns
            technical_info = self._extract_technical_patterns(text, doc_type)
            
            # Calculate processing confidence
            confidence = self._calculate_confidence(text, entities)
            
            return {
                'entities': entities,
                'embeddings': embeddings,
                'technical_info': technical_info,
                'confidence': confidence,
                'word_count': len(text.split()),
                'char_count': len(text)
            }
            
        except Exception as e:
            logger.error(f"NLP processing failed: {e}")
            return {'entities': [], 'embeddings': [], 'confidence': 0.0}
    
    def _extract_entities(self, text: str) -> List[Dict[str, Any]]:
        """Extract named entities from text"""
        if not self.nlp:
            return []
        
        doc = self.nlp(text)
        entities = []
        
        for ent in doc.ents:
            entities.append({
                'text': ent.text,
                'label': ent.label_,
                'start': ent.start_char,
                'end': ent.end_char,
                'confidence': 0.9  # spaCy doesn't provide confidence scores
            })
        
        return entities
    
    def _generate_embeddings(self, text: str) -> List[float]:
        """Generate sentence embeddings"""
        if not self.sentence_model:
            return []
        
        try:
            # Split text into chunks if too long
            max_length = 512
            if len(text) > max_length:
                chunks = [text[i:i+max_length] for i in range(0, len(text), max_length)]
                embeddings = []
                for chunk in chunks:
                    chunk_embedding = self.sentence_model.encode(chunk)
                    embeddings.extend(chunk_embedding.tolist())
                return embeddings
            else:
                embedding = self.sentence_model.encode(text)
                return embedding.tolist()
                
        except Exception as e:
            logger.error(f"Embedding generation failed: {e}")
            return []
    
    def _extract_technical_patterns(self, text: str, doc_type: str) -> Dict[str, Any]:
        """Extract technical patterns specific to industrial/factory documentation"""
        patterns = {
            'procedures': [],
            'equipment': [],
            'safety_warnings': [],
            'measurements': [],
            'schedules': []
        }
        
        # Define regex patterns for common technical elements
        import re
        
        # Equipment patterns
        equipment_pattern = r'\b(?:motor|pump|valve|sensor|controller|switch|relay|actuator|conveyor|robot|press|lathe|mill)\b'
        patterns['equipment'] = re.findall(equipment_pattern, text, re.IGNORECASE)
        
        # Safety patterns
        safety_pattern = r'\b(?:warning|caution|danger|hazard|safety|protective|emergency|lockout|tagout)\b'
        patterns['safety_warnings'] = re.findall(safety_pattern, text, re.IGNORECASE)
        
        # Measurement patterns
        measurement_pattern = r'\d+\.?\d*\s*(?:mm|cm|m|in|ft|kg|lb|psi|bar|°C|°F|rpm|hz|amp|volt|watt)'
        patterns['measurements'] = re.findall(measurement_pattern, text, re.IGNORECASE)
        
        return patterns
    
    def _calculate_confidence(self, text: str, entities: List[Dict]) -> float:
        """Calculate processing confidence based on text quality and entity extraction"""
        if not text.strip():
            return 0.0
        
        # Base confidence on text length and readability
        word_count = len(text.split())
        entity_count = len(entities)
        
        # Simple confidence calculation
        confidence = min(0.9, 0.5 + (entity_count / max(word_count / 10, 1)))
        
        return confidence

class DocumentPreprocessor:
    """Main orchestrator for document preprocessing pipeline"""
    
    def __init__(self, output_dir: str = "./processed_documents"):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(exist_ok=True)
        
        # Initialize processors
        self.format_detector = FormatDetector()
        self.text_extractor = TextExtractor()
        self.media_processor = MediaProcessor()
        self.nlp_processor = NLPProcessor()
        
        # Setup database
        self.db_path = self.output_dir / "processed_documents.db"
        self._setup_database()
        
        # Threading for parallel processing
        self.thread_executor = ThreadPoolExecutor(max_workers=4)
        
    def _setup_database(self):
        """Initialize SQLite database for storing processed documents"""
        self.conn = sqlite3.connect(str(self.db_path), check_same_thread=False)
        self.conn.execute('''
            CREATE TABLE IF NOT EXISTS processed_documents (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                file_path TEXT UNIQUE,
                file_hash TEXT,
                file_type TEXT,
                extracted_text TEXT,
                metadata TEXT,
                entities TEXT,
                embeddings TEXT,
                processing_timestamp TEXT,
                confidence_score REAL
            )
        ''')
        self.conn.commit()
    
    def _calculate_file_hash(self, file_path: str) -> str:
        """Calculate SHA-256 hash of file for deduplication"""
        hash_sha256 = hashlib.sha256()
        with open(file_path, "rb") as f:
            for chunk in iter(lambda: f.read(4096), b""):
                hash_sha256.update(chunk)
        return hash_sha256.hexdigest()
    
    def _is_already_processed(self, file_path: str, file_hash: str) -> bool:
        """Check if file has already been processed"""
        cursor = self.conn.execute(
            "SELECT id FROM processed_documents WHERE file_path = ? AND file_hash = ?",
            (file_path, file_hash)
        )
        return cursor.fetchone() is not None
    
    def process_single_file(self, file_path: str, force_reprocess: bool = False) -> Optional[ProcessedDocument]:
        """Process a single file through the complete pipeline"""
        try:
            file_path = str(Path(file_path).resolve())
            
            # Calculate file hash for deduplication
            file_hash = self._calculate_file_hash(file_path)
            
            # Check if already processed
            if not force_reprocess and self._is_already_processed(file_path, file_hash):
                logger.info(f"File already processed: {file_path}")
                return None
            
            # Detect format
            format_category, file_ext = self.format_detector.detect_format(file_path)
            
            if not self.format_detector.is_supported(file_path):
                logger.warning(f"Unsupported file format: {file_path}")
                return None
            
            logger.info(f"Processing {format_category} file: {file_path}")
            
            # Extract content based on format
            extracted_data = self._extract_content(file_path, format_category)
            
            if not extracted_data.get('text'):
                logger.warning(f"No text extracted from: {file_path}")
                return None
            
            # Process with NLP
            nlp_results = self.nlp_processor.process_text(
                extracted_data['text'], 
                format_category
            )
            
            # Create processed document
            processed_doc = ProcessedDocument(
                file_path=file_path,
                file_type=format_category,
                file_hash=file_hash,
                extracted_text=extracted_data['text'],
                metadata=extracted_data.get('metadata', {}),
                entities=nlp_results.get('entities', []),
                embeddings=nlp_results.get('embeddings', []),
                processing_timestamp=datetime.now(),
                confidence_score=nlp_results.get('confidence', 0.0),
                extracted_media=extracted_data.get('media', [])
            )
            
            # Store in database
            self._store_processed_document(processed_doc)
            
            # Save extracted text to file
            self._save_text_file(processed_doc)
            
            logger.info(f"Successfully processed: {file_path}")
            return processed_doc
            
        except Exception as e:
            logger.error(f"Failed to process {file_path}: {e}")
            return None
    
    def _extract_content(self, file_path: str, format_category: str) -> Dict[str, Any]:
        """Route to appropriate extractor based on format"""
        if format_category == 'documents':
            if file_path.lower().endswith('.pdf'):
                return self.text_extractor.extract_from_pdf(file_path)
            elif file_path.lower().endswith(('.docx', '.doc')):
                return self.text_extractor.extract_from_docx(file_path)
            else:
                # Plain text files
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    return {'text': f.read(), 'metadata': {}}
        
        elif format_category == 'images':
            return self.text_extractor.extract_from_image(file_path)
        
        elif format_category == 'videos':
            return self.media_processor.extract_from_video(file_path)
        
        elif format_category == 'audio':
            return self.media_processor.extract_from_audio(file_path)
        
        elif format_category == 'spreadsheets':
            # Handle Excel/CSV files
            try:
                if file_path.lower().endswith('.csv'):
                    df = pd.read_csv(file_path)
                else:
                    df = pd.read_excel(file_path)
                
                text_content = df.to_string()
                return {
                    'text': text_content,
                    'metadata': {
                        'shape': df.shape,
                        'columns': df.columns.tolist()
                    }
                }
            except Exception as e:
                logger.error(f"Spreadsheet processing failed: {e}")
                return {'text': '', 'metadata': {}}
        
        return {'text': '', 'metadata': {}}
    
    def _store_processed_document(self, doc: ProcessedDocument):
        """Store processed document in database"""
        self.conn.execute('''
            INSERT OR REPLACE INTO processed_documents 
            (file_path, file_hash, file_type, extracted_text, metadata, entities, embeddings, processing_timestamp, confidence_score)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
        ''', (
            doc.file_path,
            doc.file_hash,
            doc.file_type,
            doc.extracted_text,
            json.dumps(doc.metadata),
            json.dumps(doc.entities),
            json.dumps(doc.embeddings),
            doc.processing_timestamp.isoformat(),
            doc.confidence_score
        ))
        self.conn.commit()
    
    def _save_text_file(self, doc: ProcessedDocument):
        """Save extracted text to individual file"""
        output_file = self.output_dir / f"{Path(doc.file_path).stem}_{doc.file_hash[:8]}.txt"
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write(f"Source: {doc.file_path}\n")
            f.write(f"Type: {doc.file_type}\n")
            f.write(f"Processed: {doc.processing_timestamp}\n")
            f.write(f"Confidence: {doc.confidence_score:.2f}\n")
            f.write("-" * 50 + "\n\n")
            f.write(doc.extracted_text)
    
    def process_directory(self, directory_path: str, recursive: bool = True) -> List[ProcessedDocument]:
        """Process all supported files in a directory"""
        directory = Path(directory_path)
        if not directory.exists():
            logger.error(f"Directory not found: {directory_path}")
            return []
        
        # Find all files
        if recursive:
            files = [f for f in directory.rglob('*') if f.is_file()]
        else:
            files = [f for f in directory.iterdir() if f.is_file()]
        
        # Filter supported files
        supported_files = [f for f in files if self.format_detector.is_supported(str(f))]
        
        logger.info(f"Found {len(supported_files)} supported files to process")
        
        # Process files in parallel
        processed_docs = []
        futures = []
        
        for file_path in supported_files:
            future = self.thread_executor.submit(self.process_single_file, str(file_path))
            futures.append(future)
        
        # Collect results
        for future in futures:
            try:
                result = future.result(timeout=300)  # 5 minute timeout per file
                if result:
                    processed_docs.append(result)
            except Exception as e:
                logger.error(f"Processing failed: {e}")
        
        return processed_docs
    
    def get_processing_stats(self) -> Dict[str, Any]:
        """Get processing statistics"""
        cursor = self.conn.execute('''
            SELECT 
                file_type,
                COUNT(*) as count,
                AVG(confidence_score) as avg_confidence,
                MIN(processing_timestamp) as first_processed,
                MAX(processing_timestamp) as last_processed
            FROM processed_documents 
            GROUP BY file_type
        ''')
        
        stats = {}
        for row in cursor.fetchall():
            stats[row[0]] = {
                'count': row[1],
                'avg_confidence': row[2],
                'first_processed': row[3],
                'last_processed': row[4]
            }
        
        return stats
    
    def search_documents(self, query: str, limit: int = 10) -> List[Dict[str, Any]]:
        """Simple text search across processed documents"""
        cursor = self.conn.execute('''
            SELECT file_path, file_type, extracted_text, confidence_score
            FROM processed_documents
            WHERE extracted_text LIKE ?
            ORDER BY confidence_score DESC
            LIMIT ?
        ''', (f'%{query}%', limit))
        
        results = []
        for row in cursor.fetchall():
            results.append({
                'file_path': row[0],
                'file_type': row[1],
                'preview': row[2][:200] + '...' if len(row[2]) > 200 else row[2],
                'confidence': row[3]
            })
        
        return results
    
    def cleanup(self):
        """Clean up resources"""
        if hasattr(self, 'conn'):
            self.conn.close()
        self.thread_executor.shutdown(wait=True)

# Example usage and testing
if __name__ == "__main__":
    # Initialize preprocessor
    preprocessor = DocumentPreprocessor(output_dir="./processed_factory_docs")
    
    try:
        # Example: Process a single file
        # result = preprocessor.process_single_file("./sample_manual.pdf")
        
        # Example: Process entire directory
        # results = preprocessor.process_directory("./factory_documents", recursive=True)
        
        # Example: Get processing statistics
        stats = preprocessor.get_processing_stats()
        print("Processing Statistics:")
        for file_type, type_stats in stats.items():
            print(f"  {file_type}: {type_stats['count']} files, avg confidence: {type_stats['avg_confidence']:.2f}")
        
        # Example: Search documents
        # search_results = preprocessor.search_documents("safety procedure")
        # print(f"Found {len(search_results)} documents containing 'safety procedure'")
        
    finally:
        preprocessor.cleanup()